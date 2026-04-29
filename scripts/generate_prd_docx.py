from docx import Document
from docx.shared import Pt
from typing import Optional

KEY_TERMS = [
    "ACCESS",
    "QUALIFIED",
    "EVALUATION",
    "COMMITTED",
    "CLOSED",
    "LOST",
    "ACTIVE",
    "EXPIRED",
    "ON_TRACK",
    "AT_RISK",
    "STALE",
    "CRITICAL",
    "ATTENTION",
    "ADMIN",
    "MANAGER",
    "REP",
]


def add_runs_with_bold(paragraph, text: str) -> None:
    parts = text.split("**")
    for index, part in enumerate(parts):
        if not part:
            continue
        explicit_bold = index % 2 == 1
        cursor = 0
        while cursor < len(part):
            matches = [
                (part.find(term, cursor), term)
                for term in KEY_TERMS
                if part.find(term, cursor) >= 0
            ]
            if not matches:
                run = paragraph.add_run(part[cursor:])
                run.bold = explicit_bold
                break
            start, term = min(matches, key=lambda item: item[0])
            if start > cursor:
                run = paragraph.add_run(part[cursor:start])
                run.bold = explicit_bold
            run = paragraph.add_run(term)
            run.bold = True
            cursor = start + len(term)


def add_paragraph(doc: Document, text: str, style: Optional[str] = None):
    paragraph = doc.add_paragraph(style=style)
    add_runs_with_bold(paragraph, text)
    paragraph.paragraph_format.space_after = Pt(4)
    if style in {"Heading 1", "Heading 2"}:
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def main() -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    with open("PRD.md", "r", encoding="utf-8") as prd:
        for line in prd:
            text = line.rstrip()

            if not text:
                doc.add_paragraph()
                continue

            if text.startswith("# "):
                add_paragraph(doc, text[2:], "Title")
                continue

            if text.startswith("## "):
                add_paragraph(doc, text[3:], "Heading 1")
                continue

            if text.startswith("### "):
                add_paragraph(doc, text[4:], "Heading 2")
                continue

            stripped = text.lstrip()
            indent = len(text) - len(stripped)
            if stripped.startswith("- "):
                paragraph = add_paragraph(doc, stripped[2:], "List Bullet")
                paragraph.paragraph_format.left_indent = Pt(18 + indent * 3)
                continue

            add_paragraph(doc, text)

    doc.save("PRD.docx")


if __name__ == "__main__":
    main()
